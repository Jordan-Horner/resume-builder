from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from reportlab.pdfgen import canvas

from resume_builder import source_import


def manifest(vault: Path) -> dict:
    return json.loads((vault / "sources" / "manifest.json").read_text(encoding="utf-8"))


def snapshot(vault: Path, entry: dict) -> str:
    return (vault / entry["snapshot"]).read_text(encoding="utf-8")


def test_markdown_import_initializes_new_vault(tmp_path: Path, run_main) -> None:
    source = tmp_path / "resume.md"
    source.write_text("# Experience\n\nResolved priority incidents.\n", encoding="utf-8")
    vault = tmp_path / "vault"

    assert run_main(source_import.main, "--vault-root", vault, "--apply", source) == 0

    assert json.loads((vault / "vault.json").read_text())["schema_version"] == 2
    entry = manifest(vault)["sources"][0]
    assert entry["format"] == "md"
    assert "Resolved priority incidents." in snapshot(vault, entry)


def test_exact_duplicates_become_aliases(tmp_path: Path, run_main) -> None:
    sources = tmp_path / "sources"
    sources.mkdir()
    (sources / "resume-a.md").write_text("Same career fact\n", encoding="utf-8")
    (sources / "resume-b.md").write_text("Same career fact\n", encoding="utf-8")
    vault = tmp_path / "vault"

    assert run_main(source_import.main, "--vault-root", vault, "--apply", sources) == 0

    entries = manifest(vault)["sources"]
    assert len(entries) == 1
    assert entries[0]["filenames"] == ["resume-a.md", "resume-b.md"]


def test_exclusion_never_removes_registered_source(tmp_path: Path, run_main) -> None:
    sources = tmp_path / "sources"
    sources.mkdir()
    (sources / "private-notes.md").write_text("Do not import\n", encoding="utf-8")
    vault = tmp_path / "vault"
    assert run_main(source_import.main, "--vault-root", vault, "--apply", sources) == 0
    original = manifest(vault)["sources"][0]
    original_snapshot = vault / original["snapshot"]

    assert (
        run_main(
            source_import.main,
            "--vault-root",
            vault,
            "--exclude",
            "private-*",
            "--apply",
            sources,
        )
        == 0
    )

    assert manifest(vault)["sources"] == [original]
    assert original_snapshot.is_file()


def test_empty_document_is_registered_and_flagged(tmp_path: Path, run_main) -> None:
    source = tmp_path / "empty.md"
    source.write_text(" \n", encoding="utf-8")
    vault = tmp_path / "vault"

    assert run_main(source_import.main, "--vault-root", vault, "--apply", source) == 0

    entry = manifest(vault)["sources"][0]
    assert entry["extraction_status"] == "empty"
    assert entry["extracted_characters"] == 0


def test_docx_text_is_extracted(tmp_path: Path, run_main) -> None:
    source = tmp_path / "resume.docx"
    document = Document()
    document.add_paragraph("Led incident response across support teams.")
    document.save(source)
    vault = tmp_path / "vault"

    assert run_main(source_import.main, "--vault-root", vault, "--apply", source) == 0

    entry = manifest(vault)["sources"][0]
    assert "Led incident response" in snapshot(vault, entry)


def test_docx_text_boxes_are_extracted(tmp_path: Path, run_main) -> None:
    source = tmp_path / "text-box-resume.docx"
    document = Document()
    run = document.add_paragraph().add_run()
    run._r.append(
        parse_xml(
            f"<w:txbxContent {nsdecls('w')}><w:p><w:r><w:t>"
            "Owned incident communications and follow-up actions."
            "</w:t></w:r></w:p></w:txbxContent>"
        )
    )
    document.save(source)
    vault = tmp_path / "vault"

    assert Document(source).paragraphs[0].text == ""
    assert run_main(source_import.main, "--vault-root", vault, "--apply", source) == 0

    entry = manifest(vault)["sources"][0]
    assert "Owned incident communications" in snapshot(vault, entry)


def test_exact_duplicate_refreshes_a_previously_empty_extraction(tmp_path: Path, run_main) -> None:
    source = tmp_path / "resume.docx"
    document = Document()
    document.add_paragraph("Coordinated incident response across support teams.")
    document.save(source)
    vault = tmp_path / "vault"

    assert run_main(source_import.main, "--vault-root", vault, "--apply", source) == 0
    before = manifest(vault)
    entry = before["sources"][0]
    original_imported_at = entry["imported_at"]
    empty_snapshot = source_import.snapshot_content(
        entry["id"],
        entry["filenames"][0],
        entry["sha256"],
        entry["format"],
        "",
    )
    snapshot_path = vault / entry["snapshot"]
    snapshot_path.write_text(empty_snapshot, encoding="utf-8")
    entry["snapshot_sha256"] = source_import.sha256_bytes(empty_snapshot.encode("utf-8"))
    entry["extracted_characters"] = 0
    entry["extraction_status"] = "empty"
    (vault / "sources" / "manifest.json").write_text(
        json.dumps(before, indent=2) + "\n", encoding="utf-8"
    )

    assert run_main(source_import.main, "--vault-root", vault, source) == 0
    assert "Coordinated incident response" not in snapshot_path.read_text(encoding="utf-8")

    assert run_main(source_import.main, "--vault-root", vault, "--apply", source) == 0
    refreshed = manifest(vault)["sources"][0]
    assert refreshed["id"] == entry["id"]
    assert refreshed["imported_at"] == original_imported_at
    assert refreshed["extraction_status"] == "ok"
    assert refreshed["extracted_characters"] > 0
    assert "refreshed_at" in refreshed
    assert "Coordinated incident response" in snapshot(vault, refreshed)


def test_pdf_text_is_extracted(tmp_path: Path, run_main) -> None:
    source = tmp_path / "resume.pdf"
    pdf = canvas.Canvas(str(source))
    pdf.drawString(72, 720, "Improved example operations reliability")
    pdf.save()
    vault = tmp_path / "vault"

    assert run_main(source_import.main, "--vault-root", vault, "--apply", source) == 0

    entry = manifest(vault)["sources"][0]
    assert "Improved example operations reliability" in snapshot(vault, entry)


def test_preview_does_not_initialize_or_write(tmp_path: Path, run_main) -> None:
    source = tmp_path / "resume.md"
    source.write_text("Preview only.\n", encoding="utf-8")
    vault = tmp_path / "vault"

    assert run_main(source_import.main, "--vault-root", vault, source) == 0

    assert not vault.exists()


def test_configured_paths_are_honored(tmp_path: Path, run_main) -> None:
    source = tmp_path / "resume.md"
    source.write_text("Configured evidence.\n", encoding="utf-8")
    vault = tmp_path / "vault"
    vault.mkdir()
    (vault / "vault.json").write_text(
        json.dumps(
            {
                "schema_version": 2,
                "facts_path": "career-facts",
                "employment_path": "jobs",
                "sources_manifest": "evidence/manifest.json",
            }
        ),
        encoding="utf-8",
    )

    assert run_main(source_import.main, "--vault-root", vault, "--apply", source) == 0

    assert (vault / "evidence" / "manifest.json").is_file()
    assert (vault / "career-facts").is_dir()
    assert (vault / "jobs").is_dir()
    assert not (vault / "sources" / "manifest.json").exists()


def test_manifest_snapshot_cannot_escape_vault(tmp_path: Path, run_main) -> None:
    source = tmp_path / "resume.md"
    source.write_text("Evidence.\n", encoding="utf-8")
    vault = tmp_path / "vault"
    (vault / "sources").mkdir(parents=True)
    (vault / "vault.json").write_text(
        json.dumps(
            {
                "schema_version": 2,
                "facts_path": "facts",
                "employment_path": "employment",
                "sources_manifest": "sources/manifest.json",
            }
        ),
        encoding="utf-8",
    )
    outside = tmp_path / "outside.md"
    outside.write_text("Must survive.\n", encoding="utf-8")
    (vault / "sources" / "manifest.json").write_text(
        json.dumps(
            {
                "version": 1,
                "sources": [
                    {
                        "id": "SRC-0123456789ab",
                        "sha256": "0123456789ab" + "0" * 52,
                        "filenames": ["resume.md"],
                        "snapshot": str(outside),
                    }
                ],
            }
        ),
        encoding="utf-8",
    )

    assert run_main(source_import.main, "--vault-root", vault, "--apply", source) == 2

    assert outside.read_text() == "Must survive.\n"
